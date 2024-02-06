# -*- coding: UTF-8 -*-
import shutil
from NovelAIUser.settings import *
import random
from Config.models import *
from docx import Document
from docx.shared import Inches
from moviepy.editor import *
from django_pandas.io import read_frame
import re
import os

try:
    from apps.Utils.script.GeneralTools import *
except:
    from apps.Utils.pyd.GeneralTools import *




# 随机函数判断是否添加图片
def random_func():
    if random.random() < 0.3:
        return True
    else:
        return False


def step_3_process(num, Task, TaskEach):
    # 用户数据验证
    result_info = user_authenticate()
    print(result_info["msg"])

    # 剪映配置
    jy_json = get_JY_model()
    # print(jy_json)
    full_path = jy_json['path']
    image_speed = jy_json['image_speed']

    false = False
    true = True
    null = None

    if '用户验证通过' in result_info["msg"]:
        data = Task.objects.filter(id=num).first()
        project_title = data.cn_name
        en_name = data.en_name
        type_path = data.type
        task_id = data.id
        content = data.content
        # print(project_title, en_name, type_path, task_id)

        # print(task_id)
        data_list = TaskEach.objects.filter(task_id=task_id)

        # 随机抽一张图做封面
        # 图片路径
        img_path = path_audio = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "data_png")
        # 获取文件夹中所有图片文件的文件名
        img_files = [f for f in os.listdir(img_path) if f.endswith('.png')]
        # 从文件名列表中随机选择一个
        random_img_file = random.choice(img_files)
        # 构建完整的文件路径
        random_img_path = os.path.join(img_path, random_img_file)
        # 结果路径
        result_path = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "data_result", "封面.png")
        # 移动图片到结果路径
        shutil.copyfile(random_img_path, result_path)

        # 制作内容文档docx
        # 按照原文拆分并移除空元素
        content = content.split('\n')
        content = [element for element in content if element != ""]
        # 读取每个项目的excel
        df_temp_dict = {}
        for i in data_list:
            df_temp_dict[i.txt] = i.index
        # print(df_temp_dict)
        # 创建一个新的文档
        doc = Document()
        for each_content in content:
            # 将列表中的每个元素写入新文档的一行
            doc.add_paragraph(each_content)
            # 随机函数书否增加图片 7成不加 3成加
            if random_func():
                # 进行判断当前content所在索引图
                for key in df_temp_dict.keys():
                    if key in each_content:
                        # 获取当前图片索引
                        jpg_index = df_temp_dict[key]
                        # 添加一张图片，并指定宽度和高度
                        picture_path = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "data_png",
                                                    str(jpg_index) + ".png")
                        doc.add_picture(picture_path, width=Inches(2))
        word_save_path = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "data_result", project_title + ".docx")
        doc.save(word_save_path)

        # 制作剪映配置文件
        # 获取每个可编辑项目的音频路径
        each_audio_path = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "each_audio_wav")
        # 音频完整路径
        full_each_audio_path = os.path.join(full_path, each_audio_path)
        # print(full_each_audio_path)
        # 获取每个可编辑项目的图片途径
        each_jpg_path = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "data_png")
        # 图片完整路径
        full_each_jpg_path = os.path.join(full_path, each_jpg_path)
        # print(full_each_jpg_path)
        audio_duration_list = []
        for j in data_list:
            # 读取每个index代表的音频并计算时长
            audio = AudioFileClip(each_audio_path + "/" + str(j.index) + ".wav")
            # 计算时长
            duration = audio.duration
            # print(duration)
            audio_duration_list.append(duration)

        each_excel = read_frame(data_list)
        # 将每个音频的时长添加到表格中
        each_excel["duration"] = audio_duration_list
        each_excel["end"] = each_excel["duration"].cumsum()
        start_list = list(each_excel["end"])
        start_list.insert(0, 0)
        each_excel["start"] = start_list[:-1]
        # print(each_excel)
        # each_excel.to_excel("d.xlsx")

        # 基于现有的文件进行修改
        draft_content = os.path.join(BASE_DIR, "data_JianyingPro_Drafts", "draft_content.json")
        with open(draft_content, encoding='utf8') as f:
            data = json.loads(f.read())

        n = each_excel.index.max() + 1

        # 生成画布的 canvases
        canvases = []
        for i in range(len(each_excel)):
            canvases_dict = {
                "album_image": "",
                "blur": 0.0,
                "color": "",
                "id": "{}".format("canvases-" + str(each_excel["index"][i])),
                "image": "",
                "image_id": "",
                "image_name": "",
                "source_platform": 0,
                "team_id": "",
                "type": "canvas_color"
            }
            canvases.append(canvases_dict)
        data["materials"]["canvases"] = canvases

        # 生成画布动画 material_animations
        material_animations = []
        for i in range(len(each_excel)):
            material_animations_dict = {
                'animations': [],
                'id': '{}'.format("material_animations-" + str(each_excel["index"][i])),
                'type': 'sticker_animation'
            }
            material_animations.append(material_animations_dict)
        data["materials"]["material_animations"] = material_animations

        # 生成 速度配置 speeds
        speeds = []
        for i in range(len(each_excel)):
            speeds_dict = {
                'curve_speed': None,
                'id': '{}'.format("speeds-" + str(each_excel["index"][i])),
                'mode': 0,
                'speed': 1.0,
                'type': 'speed'
            }
            speeds.append(speeds_dict)
        data["materials"]["speeds"] = speeds

        # 生成文本texts
        # 备注：字符串暂时没弄
        texts = []
        for i in range(len(each_excel)):
            texts_dict = {
                'add_type': 0,
                'alignment': 1,
                'background_alpha': 1.0,
                'background_color': '',
                'background_height': 1.0,
                'background_horizontal_offset': 0.0,
                'background_round_radius': 0.0,
                'background_vertical_offset': 0.0,
                'background_width': 1.0,
                'bold_width': 0.0,
                'border_color': '#ffffff',
                'border_width': 0.08,
                'check_flag': 15,
                'content': '<outline color=(1,1,1,1) width=0.08><size=11><color=(0,0,0,1)><font id="" path="D:/MyTools/JianyingPro/3.4.1.9179/Resources/Font/SystemFont/zh-hans.ttf">[{}]</font></color></size></outline>'.format(
                    each_excel["txt"][i]),
                'font_category_id': '',
                'font_category_name': '',
                'font_id': '',
                'font_name': '',
                'font_path': 'D:/MyTools/JianyingPro/3.4.1.9179/Resources/Font/SystemFont/zh-hans.ttf',
                'font_resource_id': '',
                'font_size': 11.0,
                'font_title': 'none',
                'font_url': '',
                'fonts': [],
                'global_alpha': 1.0,
                'has_shadow': False,
                'id': '{}'.format("texts-" + str(i)),
                'initial_scale': 1.0,
                'is_rich_text': False,
                'italic_degree': 0,
                'ktv_color': '',
                'layer_weight': 1,
                'letter_spacing': 0.0,
                'line_spacing': 0.02,
                'recognize_type': 0,
                'shadow_alpha': 0.0,
                'shadow_angle': -45.0,
                'shadow_color': '#000000',
                'shadow_distance': 8.0,
                'shadow_point': {'x': 1.0182337649086284, 'y': -1.0182337649086284},
                'shadow_smoothing': 0.99,
                'shape_clip_x': False,
                'shape_clip_y': False,
                'style_name': '黑字白边',
                'sub_type': 0,
                'text_alpha': 1.0,
                'text_color': '#000000',
                'text_size': 30,
                'text_to_audio_ids': [],
                'type': 'text',
                'typesetting': 0,
                'underline': False,
                'underline_offset': 0.22,
                'underline_width': 0.05,
                'use_effect_default_color': False
            }
            texts.append(texts_dict)
        data["materials"]["texts"] = texts

        # 视频轨道video_trackings
        video_trackings = []
        for i in range(len(each_excel)):
            video_trackings_dict = {
                'config': {'center_x': 0.0,
                           'center_y': 0.0,
                           'height': 0.0,
                           'rotation': 0.0,
                           'width': 0.0},
                'enable_scale': False,
                'enable_video_tracking': True,
                'id': '{}'.format("video_trackings-" + str(each_excel["index"][i])),
                'map_path': '',
                'result_path': '',
                'tracker_type': 0,
                'trackers': [],
                'tracking_time_range': 0,
                'type': 'video_tracking',
                'version': ''
            }
            video_trackings.append(video_trackings_dict)
        data["materials"]["video_trackings"] = video_trackings

        # 视频 videos
        videos = []
        for i in range(len(each_excel)):
            videos_dict = {
                'audio_fade': None,
                'cartoon_path': '',
                'category_id': '',
                'category_name': 'local',
                'check_flag': 30719,
                'crop': {'lower_left_x': 0.0,
                         'lower_left_y': 1.0,
                         'lower_right_x': 1.0,
                         'lower_right_y': 1.0,
                         'upper_left_x': 0.0,
                         'upper_left_y': 0.0,
                         'upper_right_x': 1.0,
                         'upper_right_y': 0.0},
                'crop_ratio': 'free',
                'crop_scale': 1.0,
                'duration': 10800000000,
                'extra_type_option': 0,
                'formula_id': '',
                "freeze": null,
                'gameplay': null,
                'has_audio': false,
                'height': 1080,
                'id': 'videos-{}'.format(each_excel["index"][i]),
                'intensifies_audio_path': '',
                'intensifies_path': '',
                'is_unified_beauty_mode': false,
                'material_id': '',
                'material_name': '{}.png'.format(each_excel["index"][i]),
                'material_url': '',
                'matting': {
                    "flag": 0,
                    "has_use_quick_brush": false,
                    "has_use_quick_eraser": false,
                    "interactiveTime": [],
                    "path": "",
                    "strokes": []
                },
                "media_path": "",
                "object_locked": null,
                'path': '{}\\{}.png'.format(full_each_jpg_path, str(each_excel["index"][i])),
                "picture_from": "none",
                "picture_set_category_id": "",
                "picture_set_category_name": "",
                "request_id": "",
                'reverse_intensifies_path': '',
                'reverse_path': '',
                'source_platform': 0,
                'stable': null,
                'type': 'photo',
                'video_algorithm': {
                    "algorithms": [],
                    "deflicker": null,
                    "motion_blur_config": null,
                    "noise_reduction": null,
                    "path": "",
                    "time_range": null
                },
                'width': 1920
            }
            videos.append(videos_dict)
        data["materials"]["videos"] = videos

        # 音频 aduios
        audios = []
        for i in range(len(each_excel)):
            audio_dict = {
                "app_id": 0,
                "category_id": "",
                "category_name": "local",
                "check_flag": 1,
                "duration": 5850000,
                "effect_id": "",
                "formula_id": "",
                "id": 'audios-{}'.format(each_excel["index"][i]),
                "intensifies_path": "",
                "local_material_id": "",
                "music_id": "",
                "name": "{}.wav".format(each_excel["index"][i]),
                'path': '{}\\{}.wav'.format(full_each_audio_path, str(each_excel["index"][i])),
                "resource_id": "",
                "source_platform": 0,
                "team_id": "",
                "text_id": "",
                "tone_category_id": "",
                "tone_category_name": "",
                "tone_effect_id": "",
                "tone_effect_name": "",
                "tone_speaker": "",
                "tone_type": "",
                "type": "extract_music",
                "video_id": "",
                "wave_points": []
            }
            audios.append(audio_dict)
        data["materials"]["audios"] = audios

        # tracks 轨道主线

        # source_timerange,target_timerange 需要再定义

        # 关键帧参数组合
        def key_group(common_keyframes_speed, num):
            if jy_json['keyframe_direction'] == "四方向":
                clip_group_list = [
                    # 从左到右
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": -common_keyframes_speed,
                                "y": 0.0
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从右到左
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": common_keyframes_speed,
                                "y": 0.0
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从上到下
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": 0.0,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从下到上
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": 0.0,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                ]
            elif jy_json['keyframe_direction'] == "八方向":
                clip_group_list = [
                    # 左上到右下
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": -common_keyframes_speed,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 左下到右上
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": -common_keyframes_speed,
                                "y": -common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 右下到左上
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": common_keyframes_speed,
                                "y": -common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 右上到左下
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": common_keyframes_speed,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从左到右
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": -common_keyframes_speed,
                                "y": 0.0
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从右到左
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": common_keyframes_speed,
                                "y": 0.0
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从上到下
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": 0.0,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从下到上
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": 0.0,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                ]
            elif jy_json['keyframe_direction'] == "上下方向":
                clip_group_list = [
                    # 从上到下
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": 0.0,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从下到上
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": 0.0,
                                "y": common_keyframes_speed
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                ]
            elif jy_json['keyframe_direction'] == "左右方向":
                clip_group_list = [
                    # 从左到右
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": -common_keyframes_speed,
                                "y": 0.0
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                    # 从右到左
                    {
                        "clip": {
                            "alpha": 1.0,
                            "flip": {
                                "horizontal": false,
                                "vertical": false
                            },
                            "rotation": 0.0,
                            "scale": {
                                "x": 1.0 + common_keyframes_speed,
                                "y": 1.0 + common_keyframes_speed
                            },
                            "transform": {
                                "x": common_keyframes_speed,
                                "y": 0.0
                            }},
                        "common_keyframes": [
                            {
                                "id": "common_keyframes-x-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            -common_keyframes_speed
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-x2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            common_keyframes_speed
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionX"
                            },
                            {
                                "id": "common_keyframes-y-{}".format(each_excel["index"][num]),
                                "keyframe_list": [
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y1-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": 0,  # 关键帧开始时间
                                        "values": [
                                            0.0
                                        ]
                                    },
                                    {
                                        "curveType": "Line",
                                        "graphID": "",
                                        "id": "common_keyframes-y2-{}".format(each_excel["index"][num]),
                                        "left_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "right_control": {
                                            "x": 0.0,
                                            "y": 0.0
                                        },
                                        "time_offset": int(each_excel["duration"][num] * 1000000),  # 关键帧持续时间
                                        "values": [
                                            0.0
                                        ]
                                    }
                                ],
                                "property_type": "KFTypePositionY"
                            }
                        ],
                    },
                ]
            # 随机返回一个结果
            clip_ = random.choice(clip_group_list)
            return clip_

        # 图片&视频主轴
        tracks_0_segments = []
        # "clip": clip_group_list[0]["clip"],
        # "common_keyframes": clip_group_list[0]["common_keyframes"],
        # print(clip_group_list[0]["clip"])
        # print(clip_group_list[0]["common_keyframes"])
        for i in range(len(each_excel)):
            # clip_group_list = key_group(common_keyframes_speed, i)
            # 关键帧移动倍数
            # print(image_speed)
            common_keyframes_speed = ((each_excel["duration"][i]) * 0.5) / (each_excel["duration"][i]) * (image_speed / 100)

            tracks_0_segments_dict = {
                "cartoon": false,
                "clip": key_group(common_keyframes_speed, i)["clip"],
                "common_keyframes": key_group(common_keyframes_speed, i)["common_keyframes"],
                "enable_adjust": true,
                "enable_color_curves": true,
                "enable_color_wheels": true,
                "enable_lut": true,
                "enable_smart_color_adjust": false,
                "extra_material_refs": [
                    "speeds-{}".format(str(each_excel["index"][i])),
                    "canvases-{}".format(str(each_excel["index"][i])),
                    "video_trackings-{}".format(str(each_excel["index"][i]))
                ],
                "group_id": "",
                "hdr_settings": {
                    "intensity": 1.0,
                    "mode": 1,
                    "nits": 1000
                },
                'id': '{}'.format("tracks_0_segments-" + str(each_excel["index"][i])),
                "intensifies_audio": false,
                "is_tone_modify": false,
                "keyframe_refs": [],
                "last_nonzero_volume": 1.0,
                "material_id": "videos-{}".format(str(each_excel["index"][i])),
                "render_index": 0,
                "reverse": false,
                "source_timerange": {
                    "duration": int(each_excel["duration"][i] * 1000000),
                    "start": 0
                },
                "speed": 1.0,
                "target_timerange": {
                    "duration": int(each_excel["duration"][i] * 1000000),
                    "start": int(each_excel["start"][i] * 1000000)
                },
                "track_attribute": 0,
                "track_render_index": 0,
                "visible": true,
                "volume": 1.0
            }
            # print(tracks_0_segments_dict)
            tracks_0_segments.append(tracks_0_segments_dict)
        data["tracks"][0]["segments"] = tracks_0_segments

        # 音频主轴
        tracks_1_segments = []
        for i in range(n):
            tracks_1_segments_dict = {
                "cartoon": false,
                "clip": null,
                "enable_adjust": false,
                "enable_color_curves": true,
                "enable_color_wheels": true,
                "enable_lut": false,
                "enable_smart_color_adjust": false,
                "extra_material_refs": [
                    "speeds-{}".format(str(each_excel["index"][i])),
                    # "canvases-{}".format(str(each_excel["index"][i])),
                    "video_trackings-{}".format(str(each_excel["index"][i]))
                ],
                "group_id": "",
                "hdr_settings": null,
                'id': '{}'.format("tracks_1_segments-" + str(each_excel["index"][i])),
                "intensifies_audio": false,
                "is_placeholder": false,
                "is_tone_modify": false,
                "keyframe_refs": [],
                "last_nonzero_volume": 1.0,
                "material_id": "audios-{}".format(str(each_excel["index"][i])),
                "render_index": 0,
                "reverse": false,
                "source_timerange": {
                    "duration": int(each_excel["duration"][i] * 1000000),
                    "start": 0
                },
                "speed": 1.0,
                "target_timerange": {
                    "duration": int(each_excel["duration"][i] * 1000000),
                    "start": int(each_excel["start"][i] * 1000000)
                },
                "template_id": "",
                "template_scene": "default",
                "track_attribute": 0,
                "track_render_index": 0,
                "visible": true,
                "volume": 1.0
            }
            tracks_1_segments.append(tracks_1_segments_dict)
        data["tracks"][1]["segments"] = tracks_1_segments

        # 保存结果文件
        result_txt = json.dumps(data, ensure_ascii=False)
        # result_txt = result_txt.replace()

        result_path = os.path.join(Txt2VideoDir, type_path, en_name, "data_result")

        if not os.path.exists(result_path):
            os.makedirs(result_path)

        result_file = os.path.join(result_path, "draft_content.json")

        with open(result_file, 'w+', encoding="utf8") as f:
            f.write(result_txt)

        # AE配置文件

        each_audio_path = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "each_audio_wav")
        audio_list = os.listdir(each_audio_path)
        audio_list = sorted(audio_list, key=lambda x: int(re.search(r'\d+', x).group()))
        audio_path_full_list = [os.path.join(each_audio_path, i).replace("\\", "/") for i in audio_list]

        audio_duration_list = [AudioFileClip(audio_path).duration for audio_path in audio_path_full_list]
        # print(audio_path_full_list)

        png_path = os.path.join(BASE_DIR, "Txt2Video", type_path, en_name, "data_png")
        png_list = os.listdir(png_path)
        png_list = sorted(png_list, key=lambda x: int(re.search(r'\d+', x).group()))
        pnd_path_full_list = [os.path.join(png_path, i).replace("\\", "/") for i in png_list]

        script = 'var comp = app.project.items.addComp("My Comp", 1920, 1080, 1, ' + str(
            sum(audio_duration_list)) + ', 30);\n'

        for j in range(len(audio_duration_list)):
            script += f'var audio{j + 1} = comp.layers.add(app.project.importFile(new ImportOptions(File("{audio_path_full_list[j]}"))));\n'
            script += f'var image{j + 1} = comp.layers.add(app.project.importFile(new ImportOptions(File("{pnd_path_full_list[j]}"))));\n'
            script += f'audio{j + 1}.startTime = {sum(audio_duration_list[:j])};\n'
            script += f'image{j + 1}.startTime = {sum(audio_duration_list[:j])};\n'
            script += f'image{j + 1}.outPoint = {sum(audio_duration_list[:j + 1])};\n'

        result_path = os.path.join(Txt2VideoDir, type_path, en_name, "data_result")

        if not os.path.exists(result_path):
            os.makedirs(result_path)
        result_file = os.path.join(result_path, "AE_script.jsx")
        with open(result_file, "w") as file:
            file.write(script)

        print_info = [
            "生成配置文件完毕，请复制剪映的配置文件到指定位置进行二次剪辑",
            "文章类别：{}".format(type_path),
            "英文名称：{}".format(en_name),
            "中文名称：{}".format(project_title),
        ]

        print_with_border(print_info)
